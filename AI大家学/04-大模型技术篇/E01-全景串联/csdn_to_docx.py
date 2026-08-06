"""
Download CSDN blog post (text + images) and save as a formatted Word document.
"""
import os, re, sys, time, hashlib
sys.stdout.reconfigure(encoding='utf-8')
import urllib.request
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = r"D:\MyNewStart\从GPT、Deepseek到OpenClaw、Hermes一期视频打通AI底层逻辑"
HTML_FILE = os.path.join(BASE_DIR, "csdn_raw.html")
IMG_DIR = os.path.join(BASE_DIR, "csdn_images")
OUTPUT_DOCX = os.path.join(BASE_DIR, "CSDN_LLM_AgentSkill_原文.docx")

os.makedirs(IMG_DIR, exist_ok=True)

# ── 1. Parse HTML ──────────────────────────────────────────────
with open(HTML_FILE, "r", encoding="utf-8") as f:
    soup = BeautifulSoup(f.read(), "html.parser")

# Title
title_el = soup.find("h1", class_="title-article") or soup.find("h1")
title_text = title_el.get_text().strip() if title_el else "Unknown Title"
print(f"Title: {title_text}")

# Article body — CSDN nests real content in markdown_views div
article = soup.find("div", class_="markdown_views")
if not article:
    article = soup.find("div", id="article_content")
if not article:
    article = soup.find("div", class_="content_views")
print(f"Article length: {len(article.get_text())} chars")

# ── 2. Download images ─────────────────────────────────────────
img_map = {}  # url -> local path
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
    "Referer": "https://blog.csdn.net/",
}

all_imgs = article.find_all("img")
print(f"Downloading {len(all_imgs)} images...")

for i, img in enumerate(all_imgs):
    src = img.get("src", "")
    if not src or src.startswith("data:"):
        continue
    # Clean URL
    src = src.split("#")[0].split("?")[0]
    if src in img_map:
        continue
    # Generate filename
    ext = os.path.splitext(src)[-1] or ".png"
    if ext not in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".bmp"):
        ext = ".png"
    fname = f"img_{i:03d}{ext}"
    local_path = os.path.join(IMG_DIR, fname)

    try:
        req = urllib.request.Request(src, headers=headers)
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = resp.read()
            with open(local_path, "wb") as f:
                f.write(data)
        img_map[src] = local_path
        print(f"  [{i+1}/{len(all_imgs)}] OK: {fname} ({len(data)//1024}KB)")
    except Exception as e:
        print(f"  [{i+1}/{len(all_imgs)}] FAIL: {e}")
    time.sleep(0.1)

print(f"Downloaded {len(img_map)} images successfully.")

# ── 3. Build Word Document ─────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# Configure styles
style_normal = doc.styles["Normal"]
style_normal.font.name = "Microsoft YaHei"
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15

# Title
p = doc.add_heading(title_text, level=0)
for run in p.runs:
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

# Helper: add a run with formatting
def add_run(paragraph, text, bold=False, italic=False, code=False, color=None, size=None):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return run

# Helper: process inline elements (bold, code, links, etc.)
def process_inline(paragraph, element):
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip():
            paragraph.add_run(text)
        elif text:
            paragraph.add_run(text)
        return
    if not isinstance(element, Tag):
        return

    tag = element.name
    if tag in ("strong", "b"):
        run = paragraph.add_run(element.get_text())
        run.bold = True
    elif tag in ("em", "i"):
        run = paragraph.add_run(element.get_text())
        run.italic = True
    elif tag == "code":
        add_run(paragraph, element.get_text(), code=True)
    elif tag == "a":
        run = paragraph.add_run(element.get_text())
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        run.underline = True
    elif tag == "img":
        pass  # handled at block level
    elif tag == "br":
        paragraph.add_run("\n")
    elif tag == "span":
        # Recurse into span children
        for child in element.children:
            process_inline(paragraph, child)
    else:
        # Default: just add text, recurse into children
        for child in element.children:
            process_inline(paragraph, child)

# Process block-level elements
def process_element(el, heading_level=0):
    if isinstance(el, NavigableString):
        text = str(el).strip()
        if text:
            doc.add_paragraph(text)
        return
    if not isinstance(el, Tag):
        return

    tag = el.name
    classes = el.get("class", [])
    if isinstance(classes, str):
        classes = classes.split()

    # Skip hidden elements, scripts, styles, TOC, SVG
    if tag in ("script", "style", "noscript", "svg"):
        return
    if "toc" in classes:
        return

    # Headings
    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        p = doc.add_heading(el.get_text().strip(), level=level)
        return

    # Images
    if tag == "img":
        src = el.get("src", "")
        src = src.split("#")[0].split("?")[0]
        if src in img_map:
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_map[src], width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败: {e}]")
        return

    # Code blocks
    if tag == "pre":
        code_el = el.find("code")
        code_text = code_el.get_text() if code_el else el.get_text()
        # Add as styled code block
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Add shading via run
        run = p.add_run(code_text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
        return

    # Paragraphs
    if tag == "p":
        # Check if paragraph contains only an image
        imgs = el.find_all("img", recursive=False)
        if imgs:
            for img in imgs:
                process_element(img)
            return
        # Regular paragraph with inline content
        p = doc.add_paragraph()
        for child in el.children:
            process_inline(p, child)
        # Check for nested images
        for img in el.find_all("img"):
            process_element(img)
        return

    # Lists
    if tag in ("ul", "ol"):
        for li in el.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
            for child in li.children:
                process_inline(p, child)
        return

    # Tables
    if tag == "table":
        rows = el.find_all("tr")
        if rows:
            # Determine column count
            max_cols = 0
            for row in rows:
                cols = row.find_all(["td", "th"])
                max_cols = max(max_cols, len(cols))
            if max_cols > 0:
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for j, cell in enumerate(cells):
                        if j < max_cols:
                            table.rows[i].cells[j].text = cell.get_text().strip()
                return
        return

    # Blockquote
    if tag == "blockquote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(el.get_text().strip())
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        return

    # Figure / div containing images
    if tag == "figure" or (tag == "div" and el.find("img")):
        imgs = el.find_all("img")
        figcaption = el.find("figcaption")
        for img in imgs:
            process_element(img)
        if figcaption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(figcaption.get_text().strip())
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.italic = True
        return

    # For other container elements (div, section, etc.), recurse into children
    for child in el.children:
        process_element(child)

# Process all top-level children of the article
for child in article.children:
    process_element(child)

# ── 4. Save ────────────────────────────────────────────────────
doc.save(OUTPUT_DOCX)
print(f"\nDone! Word document saved: {OUTPUT_DOCX}")
print(f"   File size: {os.path.getsize(OUTPUT_DOCX) // 1024} KB")
